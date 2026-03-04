"""
Generates ECU Reflash Tracker User Manual as a Word (.docx) document.
Run from any directory — output is placed next to this script's parent folder.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour palette ─────────────────────────────────────────────────────────────
C_BRAND    = RGBColor(0x63, 0x66, 0xF1)   # indigo (primary)
C_SUCCESS  = RGBColor(0x22, 0xC5, 0x5E)   # green
C_WARN     = RGBColor(0xF5, 0x9E, 0x0B)   # amber
C_ERROR    = RGBColor(0xEF, 0x44, 0x44)   # red
C_DARK_BG  = RGBColor(0x1E, 0x21, 0x30)   # navy
C_TEXT_DIM = RGBColor(0x8B, 0x90, 0xA8)   # grey
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER   = RGBColor(0x1E, 0x21, 0x30)   # table header bg
C_ROW_ALT  = RGBColor(0xF3, 0xF4, 0xF9)   # alternating row


def set_cell_bg(cell, hex_color: str):
    """Set background shading of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_styled_table(doc, headers, rows, col_widths_cm=None):
    """Build a styled table with a dark header row and alternating body rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1E2130')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F3F4F9'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths_cm:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[j])

    return table


def add_code_block(doc, lines: list[str]):
    """Add a monospace-styled text block (simulates a code/diagram box)."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2E, 0x33, 0x48)
    # Small space after block
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_info_box(doc, text: str, color: RGBColor = None):
    """Add a highlighted note/info paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = color or C_TEXT_DIM


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = C_BRAND
        elif level == 2:
            run.font.color.rgb = C_DARK_BG
        else:
            run.font.color.rgb = RGBColor(0x44, 0x48, 0x60)
    return h


def body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


# ══════════════════════════════════════════════════════════════════════════════
def build(output_path: str):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Default paragraph font ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('⚡ ECU Reflash Tracker')
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = C_BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('User Manual')
    r2.font.size = Pt(18)
    r2.font.color.rgb = C_TEXT_DIM

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run('Version 1.0  ·  March 2026  ·  English')
    r3.font.size = Pt(11)
    r3.font.color.rgb = C_TEXT_DIM

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 0. INSTALLATION & FIRST-TIME SETUP
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '0. Installation & First-Time Setup', 1)
    body(doc,
         'This section covers everything needed to run ECU Reflash Tracker on a new machine '
         'after copying the repository.')

    heading(doc, 'Prerequisites', 2)
    body(doc,
         'Only one tool needs to be installed: Docker Desktop. '
         'It bundles Docker Engine, Docker Compose, and a GUI dashboard.')

    add_styled_table(doc,
        headers=['Operating System', 'How to install'],
        rows=[
            ['Windows 10 / 11', 'Download from https://www.docker.com/products/docker-desktop'],
            ['macOS',           'Download from https://www.docker.com/products/docker-desktop'],
            ['Ubuntu / Debian', 'sudo apt install docker.io docker-compose-plugin'],
        ],
        col_widths_cm=[4.5, 11.5]
    )

    add_info_box(doc,
        'Windows: Enable the WSL 2 backend when prompted during Docker Desktop installation. '
        'Restart the computer afterwards.',
        C_WARN)

    heading(doc, 'Minimum Hardware', 2)
    add_styled_table(doc,
        headers=['Resource', 'Minimum'],
        rows=[
            ['RAM',        '4 GB free'],
            ['Disk',       '3 GB free (Docker images + database)'],
            ['Ports free', '3000, 8000, 5432, 9000, 9001'],
        ],
        col_widths_cm=[4, 12]
    )

    heading(doc, 'Step-by-step: First Run', 2)

    body(doc, '1.  Copy or clone the repository to any folder on the new machine.')
    add_code_block(doc, [
        'ecu-reflash-tracker/',
        '  backend/',
        '  frontend/',
        '  docker-compose.yml',
        '  start-network.ps1   <- Windows network helper',
        '  ...',
    ])

    body(doc, '2.  Start all services.')
    add_code_block(doc, [
        '# Windows (PowerShell)',
        'cd ecu-reflash-tracker',
        'docker compose up --build',
        '',
        '# Linux / macOS',
        'cd ecu-reflash-tracker',
        'docker compose up --build',
    ])

    body(doc,
         'The first run downloads Docker images and builds the app (3-5 minutes depending on '
         'internet speed). Subsequent starts take ~20 seconds.')

    body(doc, '3.  Wait until all containers are healthy:')
    add_code_block(doc, [
        '  Container ecu-postgres   Healthy',
        '  Container ecu-minio      Healthy',
        '  Container ecu-backend    Started',
        '  Container ecu-frontend   Started',
    ])

    body(doc, '4.  Open the app.')
    add_styled_table(doc,
        headers=['URL', 'What it is'],
        rows=[
            ['http://localhost:3000',      'The web application'],
            ['http://localhost:8000/docs', 'API documentation (Swagger)'],
            ['http://localhost:9001',      'MinIO file storage console (optional)'],
        ],
        col_widths_cm=[6, 10]
    )

    body(doc, '5.  Log in with the default credentials.')
    add_styled_table(doc,
        headers=['Role', 'Email', 'Password'],
        rows=[
            ['Admin', 'admin@local', 'admin123'],
            ['Tech',  'tech@local',  'tech123'],
        ],
        col_widths_cm=[3, 6, 7]
    )
    add_info_box(doc, 'Change passwords after first login: Profile -> Change Password.', C_WARN)

    heading(doc, 'Make it accessible from other devices', 2)
    body(doc,
         'Run the network startup script instead of docker compose up. '
         'It auto-detects your LAN IP, patches docker-compose.yml, and rebuilds the frontend.')
    add_code_block(doc, ['.\\start-network.ps1'])
    body(doc, 'See Section 14 — Network Startup Script for full details.')

    heading(doc, 'Stopping and Restarting', 2)
    add_code_block(doc, [
        '# Stop (keeps all data)',
        'docker compose down',
        '',
        '# Start again without rebuilding',
        'docker compose up -d',
        '',
        '# Full reset — deletes all data and rebuilds from scratch',
        'docker compose down -v',
        'docker compose up --build',
    ])

    heading(doc, 'Troubleshooting', 2)
    add_styled_table(doc,
        headers=['Problem', 'Solution'],
        rows=[
            ['Port already in use',
             'Change the left side of the port in docker-compose.yml (e.g. "3001:3000")'],
            ['Backend won\'t start',
             'Run:  docker compose logs backend'],
            ['Database empty after restart',
             'Data lives in Docker volumes — survives "down" but not "down -v"'],
            ['Docker Desktop not running',
             'Open Docker Desktop and wait for the whale icon to stop animating'],
        ],
        col_widths_cm=[5, 11]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '1. Overview', 1)
    body(doc,
         'ECU Reflash Tracker is a production-floor web application for tracking ECU learning '
         'and reflashing operations across multiple stations and work sessions.')

    heading(doc, 'System Hierarchy', 2)
    add_code_block(doc, [
        'Session  ──▶  Boxes  ──▶  ECUs',
        '   │',
        '   └──▶  Stations (physical flash benches)',
    ])

    add_styled_table(doc,
        headers=['Concept', 'Description'],
        rows=[
            ['Session',  'A production run (e.g. "Line A – Q1 Batch"). Groups related boxes.'],
            ['Box',      'A physical box of ECUs to be processed. Goes through learning then flashing.'],
            ['ECU',      'A single electronic control unit inside a box.'],
            ['Station',  'A physical flash bench. Can work on multiple boxes in learning, but flashes only one at a time.'],
        ],
        col_widths_cm=[4, 12]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. ROLES & PERMISSIONS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '2. Roles & Permissions', 1)
    body(doc, 'The system has two roles: Admin and Tech.')

    heading(doc, 'Admin', 2)
    for item in [
        'Create / delete sessions',
        'Create stations and assign members',
        'Add boxes to sessions',
        'Manage users (invite, change role, delete)',
        'Save and manage session templates',
        'Download session reports (CSV)',
        'Control session lifecycle (Mark Ready, Start, Close, Reopen)',
        'Can also perform all Tech actions',
    ]:
        bullet(doc, item)

    heading(doc, 'Tech', 2)
    for item in [
        'Add boxes to sessions',
        'Operate the Station Workbench',
        'Scan ECUs into boxes',
        'Freeze inventory (lock ECU list)',
        'Flash ECUs and record results',
        'Upload attachments to ECUs',
    ]:
        bullet(doc, item)

    add_info_box(doc, '► Rule of thumb: Admins configure the operation. Techs execute it on the floor.', C_BRAND)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. LOGGING IN
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '3. Logging In', 1)
    for step in [
        '1.  Open the application URL in your browser.',
        '2.  Enter your email and password.',
        '3.  Click Login.',
    ]:
        bullet(doc, step)

    body(doc,
         'After login, your name and avatar appear in the top-right navbar on every page. '
         'Click your name to open your Profile where you can change your password and toggle the '
         'language between English (EN) and Spanish (ES).')

    heading(doc, 'Login Flow', 2)
    add_code_block(doc, [
        'User                      App                   Backend API',
        ' │  Enter email+password   │                         │',
        ' │ ──────────────────────▶ │  POST /auth/login        │',
        ' │                         │ ───────────────────────▶ │',
        ' │                         │  JWT token + profile     │',
        ' │                         │ ◀─────────────────────── │',
        ' │  Redirect to Dashboard  │                         │',
        ' │ ◀────────────────────── │                         │',
    ])
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. SESSION DASHBOARD
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '4. Session Dashboard', 1)
    body(doc,
         'The home screen. All sessions in the system are listed in a table. '
         'Click any row to open that session. Admins can click + New Session to create one.')

    heading(doc, 'Dashboard Layout', 2)
    add_code_block(doc, [
        '┌─────────────────────────────────────────────────────────┐',
        '│  ⚡ ECU Reflash        👤 John          👤 Users  Logout │',
        '├─────────────────────────────────────────────────────────┤',
        '│  Flash Sessions                          [+ New Session] │',
        '│                                                          │',
        '│  Name            SW Version   Status    Created          │',
        '│  ──────────────────────────────────────────────────────  │',
        '│  Line A – Q1     v2.5.1-PROD  active    01 Mar 2026      │',
        '│  Line B – Feb    v2.4.0       completed 14 Feb 2026      │',
        '│                                                          │',
        '│  ── 📁 Session Templates ──────────────── 2 / 10 ──     │',
        '│  ┌──────────────────┐  ┌──────────────────┐             │',
        '│  │ 📋 Line A Std    │  │ 📋 Full Batch     │             │',
        '│  │ v2.5.1-PROD      │  │ v2.4.0-QA        │             │',
        '│  │ [Use →]  [✕]     │  │ [Use →]  [✕]     │             │',
        '│  └──────────────────┘  └──────────────────┘             │',
        '└─────────────────────────────────────────────────────────┘',
    ])
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. SESSION LIFECYCLE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '5. Session Lifecycle', 1)
    body(doc,
         'Sessions follow a strict state machine controlled by admins. '
         'Techs cannot change the session status.')

    heading(doc, 'State Diagram', 2)
    add_code_block(doc, [
        '                  "Mark Ready"         "Start Session"',
        '  [Created] ──────────────────▶ ready ───────────────▶ active',
        '    draft                                                  │',
        '                                           "Close Session" │',
        '                                                           ▼',
        '                                                       completed',
        '                                      "Reopen Session"     │',
        '                                active ◀──────────────────┘',
    ])

    add_styled_table(doc,
        headers=['Status', 'Badge Colour', 'Meaning'],
        rows=[
            ['draft',     'Grey',   'Being set up; not yet started'],
            ['ready',     'Amber',  'Configured and cleared for start'],
            ['active',    'Blue',   'Production in progress'],
            ['completed', 'Green',  'All work done'],
        ],
        col_widths_cm=[3.5, 3.5, 9]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. SESSION DETAIL — BOXES TAB
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '6. Session Detail — Boxes Tab', 1)

    heading(doc, '6.1  Views', 2)
    body(doc,
         'Switch between Kanban (default, one column per box status) and Grid '
         '(compact card view) using the toggle buttons in the top right.')
    add_code_block(doc, [
        '┌─ Pending ──┐  ┌─ Learning ─┐  ┌─ In Progress ┐  ┌─ Blocked ──┐  ┌─ Completed ┐',
        '│ BOX-001    │  │ BOX-003    │  │ BOX-005      │  │ BOX-007    │  │ BOX-009    │',
        '│ BOX-002    │  │ BOX-004    │  │ BOX-006      │  │            │  │ BOX-010    │',
        '└────────────┘  └────────────┘  └──────────────┘  └────────────┘  └────────────┘',
    ])

    heading(doc, '6.2  Filtering & Sorting', 2)
    add_styled_table(doc,
        headers=['Control', 'Purpose'],
        rows=[
            ['Search box',             'Filter by box serial number'],
            ['Status dropdown',        'Show only boxes in a specific status'],
            ['Station dropdown',       'Show only boxes assigned to one station'],
            ['Has Issues checkbox',    'Show only boxes with failed/scratch ECUs'],
            ['Sort selector',          'Sort by serial, created date, completion, ECU count, or issue count'],
        ],
        col_widths_cm=[5, 11]
    )

    heading(doc, '6.3  Box Lifecycle', 2)
    add_code_block(doc, [
        '  [Added to session]',
        '        │',
        '        ▼ Station claims the box',
        '     learning',
        '        │',
        '        ▼ Inventory frozen (Freeze button)',
        '    in_progress',
        '        │',
        '        ├── All ECUs OK ──────────────▶ completed',
        '        │',
        '        └── ECU(s) failed ────────────▶ blocked',
        '                                           │',
        '                                 Rework/Scratch',
        '                                           │',
        '                                      in_progress',
    ])

    heading(doc, '6.4  Box Detail Drawer', 2)
    body(doc,
         'Click any box card to open a right-side drawer showing: summary KPIs '
         '(ECU counts, durations, station), full ECU list, and quick actions '
         '(change status, open workbench).')

    heading(doc, '6.5  Adding a Box', 2)
    add_styled_table(doc,
        headers=['Field', 'Required', 'Description'],
        rows=[
            ['Box Serial',           '✅ Yes',    'Unique identifier printed on the physical box'],
            ['Expected ECU Count',   'Optional', 'Pre-declared capacity; used for progress %'],
        ],
        col_widths_cm=[4.5, 3, 8.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. SESSION DETAIL — STATIONS TAB
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '7. Session Detail — Stations Tab', 1)

    heading(doc, '7.1  Adding a Station (Admin)', 2)
    for step in [
        '1.  Click + Add Station.',
        '2.  Enter a Station Name (e.g. "Station A").',
        '3.  Tick the members to assign from the scrollable list.',
        '4.  Optionally click a preset chip to pre-fill name and members.',
        '5.  Click Add Station.',
    ]:
        bullet(doc, step)

    heading(doc, '7.2  Managing Station Members', 2)
    body(doc, 'Click a station card to open the Station Detail Modal:')
    for item in [
        'Add members via the dropdown + "Agregar" button.',
        'Remove members with the ✕ button next to their name.',
        'Save the current configuration as a Preset (💾 button) for reuse in future sessions.',
    ]:
        bullet(doc, item)

    heading(doc, '7.3  Navigating to Workbench', 2)
    body(doc,
         'From the station card or the Station Detail Modal, click Go to Workbench → '
         'to open the station\'s operating screen.')
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. ANALYTICS TAB
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '8. Session Detail — Analytics Tab', 1)
    body(doc,
         'Switch to the Analytics tab to see live production metrics updated every 5 seconds.')

    add_styled_table(doc,
        headers=['Chart', 'What it tells you'],
        rows=[
            ['Bar: ECUs per Box',                   'Which boxes have the most ECUs — helps spot outliers'],
            ['Pie: ECU Status breakdown',            'Overall success / failed / scratch proportion'],
            ['Line: ECUs/h per Station (30 min)',    'Throughput trend per workbench over time'],
            ['Area: Cumulative ECUs flashed',        'Running total vs a target reference line'],
        ],
        col_widths_cm=[6.5, 9.5]
    )

    add_info_box(doc,
        '► Click 📥 Download Report (admin only) to export a full session CSV.',
        C_BRAND)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. STATION WORKBENCH
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '9. Station Workbench', 1)
    body(doc,
         'The Workbench is the screen technicians use on the production floor. '
         'It can be opened full-screen on a dedicated terminal at the workbench URL: '
         '/sessions/:id/workbench')

    heading(doc, '9.1  Overall Workbench Flow', 2)
    add_code_block(doc, [
        '  Open Workbench',
        '       │',
        '  Select Station ──▶ Station Dashboard',
        '       │                    │',
        '       │         ┌──────────┴────────────┐',
        '       │    Scan/Claim             Open existing',
        '       │    new box                learning box',
        '       │         └──────────┬────────────┘',
        '       │                    ▼',
        '       │             Learning Phase',
        '       │          (Scan ECUs one by one)',
        '       │                    │',
        '       │          Click 🔒 Freeze Inventory',
        '       │                    │',
        '       │             Flashing Phase',
        '       │          (Flash each ECU)',
        '       │                    │',
        '       │    ┌───────────────┴──────────────┐',
        '       │  All OK                       ECU Failed',
        '       │    │                               │',
        '       │  Green banner               Rework / Scratch',
        '       │  ✅ Finish Box                     │',
        '       └───────────────────────────────────┘',
        '            Returns to Station Dashboard',
    ])

    heading(doc, '9.2  Station Dashboard', 2)
    body(doc,
         'The central hub after selecting a station. Shows all boxes currently assigned.')
    add_code_block(doc, [
        '┌──────────────────────────────────────────────────────────┐',
        '│ 🏭 Station A                        [Change Station]    │',
        '│                                                          │',
        '│  Scan New Box                                            │',
        '│  [BOX-SERIAL___________]  [Claim]                        │',
        '│  ⚠ Finish flashing BOX-007 before starting another      │',
        '│                                                          │',
        '│  📖 LEARNING (2) ──────────────────────────────────     │',
        '│  📦 BOX-001  learning   12 ECUs          [Open →]       │',
        '│  📦 BOX-002  learning    8 ECUs          [Open →]       │',
        '│                                                          │',
        '│  ⚡ FLASHING (1) ──────────────────────────────────     │',
        '│  📦 BOX-007  in_progress  20 ECUs   [Flash Table →]     │',
        '└──────────────────────────────────────────────────────────┘',
    ])

    add_info_box(doc, '✅ Multiple boxes can be in Learning simultaneously.', C_SUCCESS)
    add_info_box(doc, '❌ Only ONE box can be in Flashing at a time per station.', C_ERROR)

    heading(doc, '9.3  Learning Phase', 2)
    add_code_block(doc, [
        '┌────────────────────────────────────────────────────────┐',
        '│ [← Back]  Learning: BOX-001   12 ECUs scanned          │',
        '│                               [🔒 Freeze Inventory]    │',
        '│                                                         │',
        '│  [Scan ECU barcode…__________] [Add]                   │',
        '│                                                         │',
        '│  #   ECU Code       Status                             │',
        '│  1   ECU-AA001      learned                            │',
        '│  2   ECU-AA002      learned                            │',
        '│  …                                                     │',
        '└────────────────────────────────────────────────────────┘',
    ])
    for step in [
        '1.  Scan (or type) each ECU barcode and press Add.',
        '2.  Once all ECUs are in, click 🔒 Freeze Inventory to lock the list.',
        '3.  Click ← Back at any time to return to the dashboard without losing progress.',
    ]:
        bullet(doc, step)

    heading(doc, '9.4  Flashing Phase', 2)
    add_code_block(doc, [
        '┌──────────────────────────────────────────────────────────────────┐',
        '│ [← Back]  Flashing: BOX-007   15/20 done                        │',
        '│                                                                  │',
        '│  ECU Code   Status/Time   Att  Result   Notes   Action          │',
        '│  ─────────────────────────────────────────────────────────────  │',
        '│  ECU-001    ✅ success  2m30s  1   —       —     —              │',
        '│  ECU-002    ⚡ flashing 01:15  1  [Succ▼] [notes] [Finish]      │',
        '│  ECU-003    ✗ failed   5m10s  2   —       —    [Rework][Scratch]│',
        '│  ECU-004    learned     —     —   —       —    [Start Flash]    │',
        '└──────────────────────────────────────────────────────────────────┘',
    ])

    add_styled_table(doc,
        headers=['Action', 'Available when', 'Description'],
        rows=[
            ['Start Flash', 'Status = learned',        'Marks ECU as flashing; starts a live timer'],
            ['Finish',      'Status = flashing',       'Records result (Success/Failed) and optional notes'],
            ['Rework',      'Status = failed',         'Puts ECU back to rework_pending for another flash attempt'],
            ['🗑 Scratch',   'Status = failed / rework_pending', 'Marks as damaged — excluded from count, box can still close'],
        ],
        col_widths_cm=[3.5, 5, 7.5]
    )

    heading(doc, '9.5  Completing a Box', 2)
    body(doc,
         'When the last ECU is finished, the flash table stays open — the tech can review '
         'everything. A green banner appears:')
    add_code_block(doc, [
        '┌─────────────────────────────────────────────────────────────────┐',
        '│ 🎉  All ECUs completed — box ready to close                    │',
        '│     📖 2 boxes in learning waiting to be frozen   [✅ Finish Box] │',
        '└─────────────────────────────────────────────────────────────────┘',
    ])
    body(doc, 'Click ✅ Finish Box to: mark the box as completed, return to the Station Dashboard, '
              'and see a toast with the count of learning boxes ready to freeze next.')

    heading(doc, '9.6  Blocked Boxes', 2)
    body(doc,
         'If any ECU is marked as failed and not reworked/scratched, the box becomes Blocked. '
         'A dedicated screen appears with two options:')
    bullet(doc, '← Back — return to the dashboard without changing anything.')
    bullet(doc, 'View Flash Table — open the table to rework the failed ECUs.')
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. ECU DETAIL & ATTACHMENTS
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '10. ECU Detail & Attachments', 1)
    body(doc,
         'Click any underlined ECU code anywhere in the app to open the ECU Detail Modal.')
    add_code_block(doc, [
        '┌──────────────────────────────────────────────────────┐',
        '│  ECU-AA001                                    [✕]   │',
        '│  ┌──────────┬──────────────┬───────────────────┐   │',
        '│  │ Details  │ Flash History│ Attachments        │   │',
        '│  └──────────┴──────────────┴───────────────────┘   │',
        '│                                                     │',
        '│  Details tab:                                       │',
        '│  Status:  success         Version:  v2.5.1          │',
        '│  Attempts: 2              Last duration: 2m 18s     │',
        '│                                                     │',
        '│  Flash History tab:                                 │',
        '│  #  Started        Duration  Result    Notes        │',
        '│  1  01/03 09:15    5m 10s    failed    cable issue  │',
        '│  2  01/03 09:24    2m 18s    success   —            │',
        '│                                                     │',
        '│  Attachments tab:                                   │',
        '│  Type: [photo ▼]  Notes: [____________]             │',
        '│  [📎 Choose file]                                   │',
        '│  ─────────────────────────────────────             │',
        '│  photo  — cable_check.jpg         [↓ Download]     │',
        '└──────────────────────────────────────────────────────┘',
    ])
    body(doc, 'To upload an attachment:')
    for step in [
        '1.  Open the Attachments tab.',
        '2.  Select a Type (photo, log, other).',
        '3.  Add optional notes.',
        '4.  Click 📎 Choose file, select the file — it uploads automatically.',
    ]:
        bullet(doc, step)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. SESSION TEMPLATES
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '11. Session Templates', 1)
    body(doc,
         'Templates let admins save and reuse session configurations '
         '(name + SW version) without re-typing.')

    heading(doc, 'How to save a template', 2)
    for step in [
        '1.  Click + New Session.',
        '2.  Fill in the Session Name and Target SW Version fields.',
        '3.  Click 💾 Save as Template (appears once both fields are filled).',
        '4.  A confirmation message appears and the template is saved.',
    ]:
        bullet(doc, step)

    heading(doc, 'How to use a template', 2)
    for step in [
        '1.  On the Session Dashboard, find the 📁 Session Templates section.',
        '2.  Click Use → on the desired template card.',
        '3.  The New Session modal opens with the fields pre-filled.',
        '4.  Click Create to create the session.',
    ]:
        bullet(doc, step)

    add_styled_table(doc,
        headers=['Rule', 'Detail'],
        rows=[
            ['Maximum templates',   '10 per browser (stored in localStorage)'],
            ['Storage location',    'Browser localStorage — per device, not shared'],
            ['Duplicates',          'Rejected if same name + same SW version already exists'],
            ['Delete a template',   'Click the ✕ button on the template card'],
        ],
        col_widths_cm=[5, 11]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. USER MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '12. User Management', 1)
    body(doc,
         'Available to admins via the 👤 Users button in the top navbar.')
    body(doc, 'From the User Management modal you can:')
    for item in [
        'Invite a new user — provide name, email, role, and a temporary password.',
        'Change a user\'s role between admin and tech.',
        'Reset a user\'s password.',
        'Delete a user from the system.',
    ]:
        bullet(doc, item)

    add_styled_table(doc,
        headers=['Role', 'Can be changed to'],
        rows=[
            ['tech',  'admin'],
            ['admin', 'tech'],
        ],
        col_widths_cm=[5, 11]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. STATUS REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, '13. Status Reference', 1)

    heading(doc, 'Session Statuses', 2)
    add_styled_table(doc,
        headers=['Status', 'Badge Colour', 'Meaning'],
        rows=[
            ['draft',     'Grey',   'Created; not yet ready for production'],
            ['ready',     'Amber',  'Configured; cleared by admin to start'],
            ['active',    'Blue',   'In production'],
            ['completed', 'Green',  'Closed'],
        ],
        col_widths_cm=[3.5, 3.5, 9]
    )
    doc.add_paragraph()

    heading(doc, 'Box Statuses', 2)
    add_styled_table(doc,
        headers=['Status', 'Meaning'],
        rows=[
            ['pending',     'Added to session, not yet claimed by a station'],
            ['learning',    'Station claimed it; ECUs being scanned'],
            ['in_progress', 'Inventory frozen; flashing in progress'],
            ['blocked',     'One or more ECUs failed and need rework'],
            ['completed',   'All ECUs done successfully'],
        ],
        col_widths_cm=[4, 12]
    )
    doc.add_paragraph()

    heading(doc, 'ECU Statuses', 2)
    add_code_block(doc, [
        '  [Scanned into box]',
        '        │',
        '        ▼ Start Flash',
        '     flashing',
        '        │',
        '        ├── result=success ─────────────▶ success ──▶ [Done]',
        '        │',
        '        └── result=failed ──────────────▶ failed',
        '                                             │',
        '                              ┌─────────────┴────────────┐',
        '                           Rework                     Scratch',
        '                              │                          │',
        '                       rework_pending              scratch ──▶ [Excluded]',
        '                              │',
        '                           Re-Flash',
        '                              │',
        '                          flashing ...',
    ])
    add_styled_table(doc,
        headers=['Status', 'Colour', 'Counts toward completion?'],
        rows=[
            ['learned',        'Purple',  '—'],
            ['flashing',       'Violet',  '—  (live timer shown)'],
            ['success',        'Green',   '✅ Yes'],
            ['failed',         'Red',     'Blocks box completion'],
            ['rework_pending', 'Amber',   '—'],
            ['scratch',        'Grey',    'Excluded (does not block)'],
        ],
        col_widths_cm=[4, 3.5, 8.5]
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # 14. NETWORK STARTUP SCRIPT
    # ══════════════════════════════════════════════════════════════════════════════
    heading(doc, '14. Network Startup Script', 1)
    body(doc,
         'start-network.ps1 is a PowerShell script that auto-detects your current LAN IP, '
         'updates docker-compose.yml, and rebuilds + restarts the stack. '
         'Run it any time you change networks (office, home, client site).')

    heading(doc, 'Basic Usage', 2)
    add_code_block(doc, [
        '# From the project root',
        '.\\start-network.ps1',
    ])

    heading(doc, 'Options', 2)
    add_styled_table(doc,
        headers=['Flag', 'Default', 'Description'],
        rows=[
            ['-IP <address>',      '(auto)',   'Force a specific IP instead of auto-detecting'],
            ['-FrontendPort <n>',  '3000',     'Change the frontend port'],
            ['-BackendPort <n>',   '8000',     'Change the backend port'],
            ['-OpenFirewall',      '(off)',     'Add Windows Firewall inbound rules (requires admin)'],
        ],
        col_widths_cm=[4.5, 2.5, 9]
    )

    heading(doc, 'Examples', 2)
    add_code_block(doc, [
        '# Force a specific IP (e.g. VPN or Ethernet)',
        '.\\start-network.ps1 -IP 10.0.0.5',
        '',
        '# Use different ports',
        '.\\start-network.ps1 -FrontendPort 8080 -BackendPort 9000',
        '',
        '# Open firewall ports (run PowerShell as Administrator)',
        '.\\start-network.ps1 -OpenFirewall',
    ])

    heading(doc, 'What it does step by step', 2)
    for step in [
        '1.  Detects LAN IP — scans active adapters, prefers Wi-Fi, skips loopback/WSL/APIPA.',
        '2.  Patches docker-compose.yml — replaces VITE_API_URL and FRONTEND_URL with the detected IP.',
        '3.  Rebuilds containers — runs: docker compose up -d --build frontend backend.',
        '4.  Prints access URLs.',
    ]:
        bullet(doc, step)

    add_code_block(doc, [
        '  Local  : http://localhost:3000',
        '  Network: http://192.168.1.27:3000',
        '  API    : http://192.168.1.27:8000',
    ])

    add_info_box(doc,
        'Note: VITE_API_URL is baked into the frontend bundle at build time. '
        'A rebuild is required whenever the IP changes — the script handles this automatically.',
        C_WARN)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run('For technical setup and deployment instructions see QUICKSTART.md.')
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = C_TEXT_DIM

    doc.save(output_path)
    print(f'✅  Saved: {output_path}')


if __name__ == '__main__':
    here = os.path.dirname(os.path.abspath(__file__))
    out  = os.path.join(here, '..', 'ECU_Reflash_Tracker_User_Manual.docx')
    build(os.path.normpath(out))
