"""Genera el manual de usuario del sistema ECU Reflash Tracker en formato Word."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Manual_ECU_Reflash_Tracker.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    return p

def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")

def info_box(doc, text, bg="E8F4FD", border_color="2196F3"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    doc.add_paragraph()

def warning_box(doc, text):
    info_box(doc, "⚠  " + text, bg="FFF3E0", border_color="FF9800")

def step_table(doc, steps):
    """steps: list of (number_str, description)"""
    table = doc.add_table(rows=len(steps), cols=2)
    table.style = "Table Grid"
    for i, (num, desc) in enumerate(steps):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        set_cell_bg(c0, "1A3A5C")
        c0.width = Cm(1.2)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(num)
        r0.bold = True
        r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r0.font.size = Pt(11)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1.text = desc
        c1.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ── Document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── PORTADA ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("⚡ ECU Reflash Tracker")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x4F, 0x8E, 0xF7)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Manual de Usuario")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x55, 0x62, 0x78)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Versión 1.0  ·  Marzo 2026")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x88, 0x92, 0xA4)

doc.add_page_break()

# ─── 1. INTRODUCCIÓN ─────────────────────────────────────────────────────────
heading(doc, "1. Introducción", 1, "1A3A5C")
body(doc,
    "ECU Reflash Tracker es un sistema web diseñado para gestionar y monitorear "
    "el proceso de reprogramación (reflashing) de unidades de control electrónico (ECUs). "
    "Permite organizar cajas de ECUs, asignarlas a estaciones de trabajo, registrar el "
    "resultado de cada operación y consultar estadísticas en tiempo real.")

doc.add_paragraph()
heading(doc, "1.1  Roles de usuario", 2)

roles_table = doc.add_table(rows=4, cols=3)
roles_table.style = "Table Grid"
headers = ["Rol", "Acceso", "Descripción"]
header_row = roles_table.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_bg(cell, "1A3A5C")
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)

roles_data = [
    ("Admin",  "Total",      "Gestión completa: sesiones, cajas, estaciones, usuarios y estadísticas."),
    ("Tech",   "Operativo",  "Puede reclamar cajas, escanear ECUs, iniciar y finalizar el flasheo."),
    ("Viewer", "Solo lectura","Puede consultar sesiones y estados pero no modificar nada."),
]
for i, (role, access, desc) in enumerate(roles_data, start=1):
    row = roles_table.rows[i]
    row.cells[0].text = role
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = access
    row.cells[2].text = desc
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

doc.add_page_break()

# ─── 2. ACCESO AL SISTEMA ────────────────────────────────────────────────────
heading(doc, "2. Acceso al Sistema", 1, "1A3A5C")
body(doc, "Abre el navegador y dirígete a:")
body(doc, "http://localhost:3000", bold=True)
doc.add_paragraph()

info_box(doc,
    "Credenciales de acceso por defecto:\n"
    "  • Admin:   admin@local  /  admin123\n"
    "  • Técnico: tech@local   /  tech123\n"
    "Cambia las contraseñas en Producción desde el panel de Usuarios.")

body(doc, "Para iniciar sesión:", bold=True)
step_table(doc, [
    ("1", "Introduce tu correo electrónico en el campo Email."),
    ("2", "Introduce tu contraseña."),
    ("3", "Haz clic en el botón Login."),
])

warning_box(doc, "Si no puedes iniciar sesión, contacta al administrador para verificar tus credenciales.")

doc.add_page_break()

# ─── 3. PANEL PRINCIPAL ──────────────────────────────────────────────────────
heading(doc, "3. Panel Principal — Flash Sessions", 1, "1A3A5C")
body(doc,
    "Al ingresar verás la lista de sesiones de flasheo. Cada sesión representa "
    "un lote de cajas a programar (por ejemplo, un turno de producción).")

heading(doc, "3.1  Crear una sesión (Admin)", 2)
step_table(doc, [
    ("1", "Haz clic en el botón + New Session en la esquina superior derecha."),
    ("2", "Ingresa el Nombre de la sesión (ej. 'Lote Marzo 2026')."),
    ("3", "Ingresa la versión de software objetivo (Target SW Version)."),
    ("4", "Haz clic en Create Session."),
])

heading(doc, "3.2  Estados de una sesión", 2)
states_table = doc.add_table(rows=4, cols=2)
states_table.style = "Table Grid"
for row_data in [("active", "Sesión en curso, se pueden agregar cajas y operar."),
                 ("completed", "Todos los procesos finalizados."),
                 ("archived", "Sesión archivada, solo lectura.")]:
    row = states_table.add_row()
    row.cells[0].text = row_data[0]
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = row_data[1]
    for c in row.cells:
        c.paragraphs[0].runs[0].font.size = Pt(10)

# Remove empty first row
states_table._tbl.remove(states_table.rows[0]._tr)
doc.add_paragraph()

doc.add_page_break()

# ─── 4. GESTIÓN DE CAJAS ────────────────────────────────────────────────────
heading(doc, "4. Gestión de Cajas (Boxes)", 1, "1A3A5C")
body(doc,
    "Una caja contiene un conjunto de ECUs que serán flasheadas juntas. "
    "Las cajas se muestran en una vista Kanban o en vista de cuadrícula.")

heading(doc, "4.1  Crear una caja (Admin)", 2)
step_table(doc, [
    ("1", "Dentro de una sesión, ve a la pestaña Boxes."),
    ("2", "Haz clic en + Add Box."),
    ("3", "Ingresa el número de serie de la caja (Box Serial)."),
    ("4", "Opcionalmente ingresa el número esperado de ECUs."),
    ("5", "Haz clic en Add Box."),
])

heading(doc, "4.2  Filtrar y buscar cajas", 2)
bullet(doc, "Usa el campo de búsqueda para filtrar por número de serie.")
bullet(doc, "Usa los botones de estado (Todos, Pending, Learning, En proceso, Blocked, Completed).")
bullet(doc, "Usa el filtro de estación para ver solo las cajas asignadas a una estación.")
bullet(doc, "Activa '⚠ Con fallas' para ver únicamente cajas con ECUs fallidas o scratched.")
bullet(doc, "Usa el selector de orden (Nombre, Fecha, ECUs, Fallas) y el botón ↑↓ para ordenar.")
doc.add_paragraph()

heading(doc, "4.3  Estados de una caja", 2)
states = [
    ("Pending",      "Recién creada, sin estación asignada."),
    ("Learning",     "En proceso de escaneo de ECUs."),
    ("In Progress",  "Las ECUs están siendo flasheadas."),
    ("Blocked",      "La caja tiene ECUs con problemas que requieren atención."),
    ("Completed",    "Todas las ECUs han sido procesadas exitosamente."),
]
t = doc.add_table(rows=len(states), cols=2)
t.style = "Table Grid"
for i, (s, d) in enumerate(states):
    t.rows[i].cells[0].text = s
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    t.rows[i].cells[1].text = d
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

info_box(doc,
    "Los administradores pueden mover cajas entre estados arrastrándolas "
    "en el tablero Kanban (drag & drop).")

doc.add_page_break()

# ─── 5. ESTACIONES DE TRABAJO ───────────────────────────────────────────────
heading(doc, "5. Estaciones de Trabajo", 1, "1A3A5C")
body(doc,
    "Las estaciones representan puestos físicos de flasheo. "
    "Cada estación tiene técnicos asignados y puede tener una caja activa a la vez.")

heading(doc, "5.1  Crear una estación (Admin)", 2)
step_table(doc, [
    ("1", "Ve a la pestaña Stations dentro de la sesión."),
    ("2", "Haz clic en + Add Station."),
    ("3", "Escribe el nombre de la estación."),
    ("4", "Selecciona los técnicos miembros."),
    ("5", "Haz clic en Add Station."),
])

heading(doc, "5.2  Agregar miembros a una estación", 2)
step_table(doc, [
    ("1", "Haz clic sobre la tarjeta de la estación."),
    ("2", "En el panel lateral, busca la sección Miembros."),
    ("3", "Selecciona el técnico en el menú desplegable."),
    ("4", "Haz clic en + Agregar."),
])

doc.add_page_break()

# ─── 6. FLUJO DE TRABAJO DEL TÉCNICO ────────────────────────────────────────
heading(doc, "6. Flujo de Trabajo del Técnico", 1, "1A3A5C")
body(doc,
    "Este es el proceso paso a paso que debe seguir cada técnico durante una sesión de flasheo.")

heading(doc, "6.1  Acceder a la estación de trabajo", 2)
step_table(doc, [
    ("1", "Ingresa al sistema con tu usuario y contraseña."),
    ("2", "Ve a la sesión activa desde el panel principal."),
    ("3", "En la pestaña Stations, haz clic en tu estación y luego en Open Workbench."),
])

heading(doc, "6.2  Escaneo de ECUs (fase Learning)", 2)
info_box(doc,
    "La fase de Learning es cuando se registran todas las ECUs de la caja "
    "antes de comenzar el flasheo. La caja debe estar asignada a tu estación.")
step_table(doc, [
    ("1", "Reclama la caja haciendo clic en Claim (se te asignará automáticamente)."),
    ("2", "Escanea el código de barras de cada ECU o ingrésalo manualmente."),
    ("3", "Verifica que el número de ECUs escaneadas sea el correcto."),
    ("4", "Cuando todas las ECUs estén escaneadas, haz clic en Freeze Inventory."),
])

warning_box(doc,
    "Una vez congelado el inventario (Freeze), no se pueden agregar ni eliminar ECUs. "
    "Solo los administradores pueden eliminar ECUs antes del congelamiento.")

heading(doc, "6.3  Proceso de Flasheo", 2)
step_table(doc, [
    ("1",  "Con el inventario congelado, selecciona una ECU de la tabla."),
    ("2",  "Haz clic en Start Flash para iniciar la programación."),
    ("3",  "El sistema registra automáticamente el tiempo de inicio."),
    ("4",  "Una vez completado el proceso en el equipo físico, haz clic en Finish."),
    ("5",  "Selecciona el resultado: Success (éxito) o Failed (fallo)."),
    ("6",  "Repite para todas las ECUs de la caja."),
])
body(doc,
    "El indicador de tiempo (⏱) muestra en tiempo real cuánto lleva la operación actual. "
    "Al finalizar, se muestra la duración total con código de color.")

heading(doc, "6.4  Gestión de fallos (Rework)", 2)
body(doc,
    "Si una ECU falla durante el flasheo, tienes dos opciones:")
bullet(doc, "Rework: Envía la ECU a reproceso. Podrás volver a flashearla usando Re-Flash.")
bullet(doc, "Scratch: Marca la ECU como descartada si ya no es posible recuperarla. "
            "La caja puede cerrarse aunque haya ECUs en scratch.")
doc.add_paragraph()

info_box(doc,
    "Una caja se completa automáticamente cuando todas sus ECUs están en estado "
    "Success o Scratch.")

doc.add_page_break()

# ─── 7. ANALÍTICAS ──────────────────────────────────────────────────────────
heading(doc, "7. Analíticas y Reportes", 1, "1A3A5C")
body(doc,
    "La pestaña Analytics dentro de cada sesión muestra estadísticas en tiempo real.")

heading(doc, "7.1  Indicadores principales (KPIs)", 2)
bullet(doc, "Total Boxes / Completed / Blocked / In Progress")
bullet(doc, "Total ECUs / Success ECUs / Failed ECUs / Scratch ECUs")
bullet(doc, "Failure Rate: porcentaje de ECUs fallidas + scratched sobre el total")
doc.add_paragraph()

heading(doc, "7.2  Gráficas disponibles", 2)
bullet(doc, "ECU Results by Box: barras por caja (éxito, fallo, scratch).")
bullet(doc, "Overall ECU Status: gráfica de pastel con la distribución global.")
bullet(doc, "ECUs por Hora por Estación: productividad de cada estación por franja horaria.")
bullet(doc, "Box KPIs: tabla detallada con métricas por caja (tasa de fallo, tiempo promedio).")

doc.add_page_break()

# ─── 8. GESTIÓN DE USUARIOS ─────────────────────────────────────────────────
heading(doc, "8. Gestión de Usuarios (Admin)", 1, "1A3A5C")
body(doc,
    "Solo los administradores pueden crear, editar y eliminar usuarios.")

heading(doc, "8.1  Crear un usuario nuevo", 2)
step_table(doc, [
    ("1", "En la barra de navegación, haz clic en 👤 Usuarios."),
    ("2", "Completa el formulario: Nombre, Email, Contraseña y Rol."),
    ("3", "Haz clic en Crear usuario."),
])

heading(doc, "8.2  Editar un usuario", 2)
step_table(doc, [
    ("1", "En el panel 👤 Usuarios, localiza al usuario en la tabla."),
    ("2", "Haz clic en el botón ✏️."),
    ("3", "Modifica los campos necesarios (deja la contraseña en blanco para no cambiarla)."),
    ("4", "Haz clic en Guardar cambios."),
])

heading(doc, "8.3  Eliminar un usuario", 2)
step_table(doc, [
    ("1", "En el panel 👤 Usuarios, haz clic en 🗑 junto al usuario."),
    ("2", "Confirma la eliminación en el diálogo."),
])

warning_box(doc,
    "No puedes eliminar tu propia cuenta. "
    "Un usuario eliminado pierde acceso inmediatamente al sistema.")

doc.add_page_break()

# ─── 9. PREGUNTAS FRECUENTES ────────────────────────────────────────────────
heading(doc, "9. Preguntas Frecuentes", 1, "1A3A5C")

faqs = [
    ("¿Qué pasa si cierro el navegador durante un flasheo?",
     "El sistema conserva el estado. Al volver a abrir el workbench, la ECU seguirá "
     "en estado Flashing. Simplemente haz clic en Finish para registrar el resultado."),
    ("¿Puedo agregar ECUs después de congelar el inventario?",
     "No. Una vez congelado, el inventario es definitivo. Solo un administrador puede "
     "eliminar ECUs antes del congelamiento si se registró alguna por error."),
    ("¿Cómo sé si la caja está completada?",
     "La caja pasa automáticamente a estado Completed cuando todas sus ECUs están en "
     "Success o Scratch. En el Kanban aparece en la columna Completada."),
    ("¿Qué significa el estado 'Scratch'?",
     "Una ECU Scratch es aquella que fue descartada definitivamente (por fallos repetidos "
     "u otros motivos). Se contabiliza en la tasa de fallo pero no bloquea el cierre de la caja."),
    ("No puedo iniciar sesión, ¿qué hago?",
     "Contacta al administrador para que verifique tu usuario o restablezca tu contraseña "
     "desde el panel 👤 Usuarios."),
    ("¿Los datos se pierden si se reinicia el servidor?",
     "No. Todos los datos se almacenan en la base de datos PostgreSQL y persisten "
     "entre reinicios del sistema."),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    r_q = p_q.add_run("P: " + q)
    r_q.bold = True
    r_q.font.size = Pt(11)
    p_a = doc.add_paragraph()
    r_a = p_a.add_run("R: " + a)
    r_a.font.size = Pt(10)
    r_a.font.italic = True
    doc.add_paragraph()

# ─── FOOTER NOTE ─────────────────────────────────────────────────────────────
doc.add_page_break()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run("ECU Reflash Tracker  ·  Manual de Usuario  ·  v1.0  —  Marzo 2026")
r_end.font.size = Pt(9)
r_end.font.color.rgb = RGBColor(0x88, 0x92, 0xA4)

doc.save(OUTPUT)
print(f"Documento generado: {OUTPUT}")
